"""
Turn an uploaded complaint document into plain text for the LLM.

Supported: .pdf (via fitz/PyMuPDF), .docx, .txt, .eml

The one rule this module follows: every failure raises FileParseError with
a message written FOR THE USER, not for a log file. The API layer catches
it and posts it into the chat as an assistant message, so a bad upload
becomes "I couldn't read that file because..." instead of a red 500 or,
worse, a silent empty extraction that produces a blank form.

Nothing here calls the LLM. Parsing is deliberately dumb and offline -
production-grade OCR is explicitly out of scope for this assignment.
"""

from __future__ import annotations

import io
import re
from email import policy
from email.parser import BytesParser
from pathlib import Path
from typing import Callable, Dict

import fitz  # PyMuPDF

from app.config import settings

# gemma2-9b-it has an 8,192-token context window, and the extracted document
# is only one part of the prompt (schema + instructions + chat history share
# it). ~12k characters is roughly 3k tokens, which leaves comfortable room.
# Truncating here is far better than a 413 from Groq mid-demo.
MAX_TEXT_CHARS = 12_000

ALLOWED_EXTENSIONS = frozenset({".pdf", ".docx", ".txt", ".eml"})


class FileParseError(Exception):
    """A failure whose message is safe and useful to show directly in the chat."""


# ---------------------------------------------------------------------------
# Validation - runs BEFORE we read or parse anything
# ---------------------------------------------------------------------------


def validate_upload(filename: str, size_bytes: int) -> str:
    """
    Check extension and size. Returns the normalised lowercase extension.

    Called before parsing so we reject a 200 MB file on its metadata rather
    than after loading it into memory.
    """
    if not filename:
        raise FileParseError("That upload had no filename, so I couldn't tell what type it was.")

    extension = Path(filename).suffix.lower()

    if extension not in ALLOWED_EXTENSIONS:
        supported = ", ".join(sorted(ALLOWED_EXTENSIONS))
        # .doc is common enough to be worth a specific hint - python-docx
        # cannot read the old binary format at all.
        if extension == ".doc":
            raise FileParseError(
                "I can't read legacy .doc files. Please re-save it as .docx or .pdf "
                "and upload again."
            )
        raise FileParseError(
            f"I can't read '{extension or 'that file type'}' files. "
            f"Supported formats are: {supported}."
        )

    if size_bytes <= 0:
        raise FileParseError(f"'{filename}' appears to be empty (0 bytes).")

    if size_bytes > settings.max_upload_bytes:
        limit_mb = settings.max_upload_bytes / (1024 * 1024)
        actual_mb = size_bytes / (1024 * 1024)
        raise FileParseError(
            f"'{filename}' is {actual_mb:.1f} MB, which exceeds the {limit_mb:.0f} MB limit."
        )

    return extension


# ---------------------------------------------------------------------------
# Per-format extractors
# ---------------------------------------------------------------------------


def _extract_pdf(data: bytes) -> str:
    """PDF text via fitz. Reads from memory - nothing is written to disk."""
    try:
        with fitz.open(stream=data, filetype="pdf") as document:
            if document.needs_pass:
                raise FileParseError(
                    "That PDF is password-protected. Please upload an unlocked copy."
                )
            pages = [page.get_text("text") for page in document]
    except FileParseError:
        raise
    except Exception as exc:  # noqa: BLE001 - fitz raises assorted low-level errors
        raise FileParseError(
            f"That PDF appears to be corrupted or unreadable ({type(exc).__name__})."
        ) from exc

    text = "\n".join(pages).strip()

    # A scanned complaint letter is a real possibility in this domain: the
    # PDF opens fine but has no text layer. Without this check the user gets
    # an inexplicably blank form.
    if not text:
        raise FileParseError(
            "I opened that PDF but found no selectable text - it looks like a scanned "
            "image. Please paste the complaint text into the chat instead."
        )
    return text


def _extract_docx(data: bytes) -> str:
    """
    DOCX text via python-docx.

    Reads paragraphs AND table cells: pharma complaint forms are very often
    laid out as a table, and paragraph-only extraction would return almost
    nothing for those.
    """
    try:
        import docx  # imported lazily so a missing optional dep can't break PDF uploads

        document = docx.Document(io.BytesIO(data))
    except ImportError as exc:
        raise FileParseError(
            "DOCX support isn't installed on the server (python-docx). "
            "Please upload a PDF or paste the text instead."
        ) from exc
    except Exception as exc:  # noqa: BLE001
        raise FileParseError(
            f"That DOCX file appears to be corrupted or unreadable ({type(exc).__name__})."
        ) from exc

    parts = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                # "Batch Number | BN-2024-0783" reads clearly to the LLM.
                parts.append(" | ".join(cells))

    text = "\n".join(parts).strip()
    if not text:
        raise FileParseError("That DOCX file contains no readable text.")
    return text


def _extract_txt(data: bytes) -> str:
    """Plain text. Tries UTF-8, then Windows-1252, then replaces bad bytes."""
    for encoding in ("utf-8", "utf-8-sig", "cp1252"):
        try:
            text = data.decode(encoding).strip()
            break
        except UnicodeDecodeError:
            continue
    else:
        # Never fail on encoding alone - a mangled character is recoverable,
        # a rejected complaint is not.
        text = data.decode("utf-8", errors="replace").strip()

    if not text:
        raise FileParseError("That text file is empty.")
    return text


def _strip_html(html: str) -> str:
    """Crude tag strip for HTML-only emails. Good enough - the LLM is tolerant."""
    without_blocks = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", html)
    without_tags = re.sub(r"(?s)<[^>]+>", " ", without_blocks)
    unescaped = (
        without_tags.replace("&nbsp;", " ")
        .replace("&amp;", "&")
        .replace("&lt;", "<")
        .replace("&gt;", ">")
        .replace("&quot;", '"')
        .replace("&#39;", "'")
    )
    return re.sub(r"[ \t]*\n\s*\n\s*", "\n\n", re.sub(r"[ \t]+", " ", unescaped)).strip()


def _extract_eml(data: bytes) -> str:
    """
    Email text, headers included.

    The headers are kept on purpose and placed first: From: carries the
    customer name and Date: carries the complaint date, so handing them to
    the LLM directly fills two form fields that the body often never states.
    """
    try:
        message = BytesParser(policy=policy.default).parsebytes(data)
    except Exception as exc:  # noqa: BLE001
        raise FileParseError(
            f"That .eml file couldn't be parsed as an email ({type(exc).__name__})."
        ) from exc

    header_lines = [
        f"{label}: {message[key]}"
        for label, key in (("From", "from"), ("To", "to"), ("Date", "date"), ("Subject", "subject"))
        if message[key]
    ]

    body = ""
    if message.is_multipart():
        plain = message.get_body(preferencelist=("plain",))
        if plain is not None:
            body = plain.get_content()
        else:
            html_part = message.get_body(preferencelist=("html",))
            if html_part is not None:
                body = _strip_html(html_part.get_content())
    else:
        content = message.get_content()
        body = _strip_html(content) if message.get_content_type() == "text/html" else content

    text = "\n".join(header_lines + ["", (body or "").strip()]).strip()
    if not text:
        raise FileParseError("That email appears to have no readable body text.")
    return text


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

_EXTRACTORS: Dict[str, Callable[[bytes], str]] = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".txt": _extract_txt,
    ".eml": _extract_eml,
}


def extract_text_from_upload(filename: str, data: bytes) -> str:
    """
    Validate, parse and normalise an uploaded document.

    Returns clean plain text ready to be handed to the LangGraph pipeline.
    Raises FileParseError - and only FileParseError - on any failure, so the
    caller has exactly one exception type to turn into a chat message.
    """
    extension = validate_upload(filename, len(data))

    text = _EXTRACTORS[extension](data)

    # Collapse the runs of blank lines that PDF extraction tends to produce.
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    if len(text) > MAX_TEXT_CHARS:
        text = (
            text[:MAX_TEXT_CHARS]
            + "\n\n[Document truncated - only the first portion was analysed.]"
        )

    return text
